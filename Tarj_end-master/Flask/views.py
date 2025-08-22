from main import app
from flask import Flask, render_template, request, send_file, redirect, url_for, session, jsonify
import re
import tempfile
import os
import fitz  #
from docx import Document
import io
from historico_utils import salvar_envio
import uuid
import base64
from docx.shared import RGBColor
import json
from fuzzywuzzy import fuzz
from pdf2image import convert_from_bytes
import pytesseract
from pyzbar.pyzbar import decode
from PIL import Image
from PIL import Image, ImageDraw
import fitz  # PyMuPDF

# pyzbar (ZBar). Se falhar, loga e segue sem QR.
try:
    from pyzbar.pyzbar import decode as zbar_decode
    _ZBAR_OK = True
except Exception as _e:
    print("DEBUG pyzbar import falhou:", _e)
    _ZBAR_OK = False

from regex_patterns import PADROES_SENSIVEIS

import re
from flask import jsonify, session
from docx import Document

# Habilita sessão para guardar dados temporários
app.secret_key = "segredo-muito-seguro"


@app.route("/",  methods=["GET", "POST"])
def homepage():
    
     return render_template("index.html")
    

def copiar_e_tarjar(original_doc, padroes):
    novo_doc = Document()

    for par in original_doc.paragraphs:
        texto = par.text
        for nome, regex in padroes.items():
            texto = re.sub(regex, lambda m: "█" * len(m.group()), texto)

        novo_doc.add_paragraph(texto)

    return novo_doc

# ----------------------------------------------------------------------------------- Padrões para DOCX ----------------------------------------------------------------------------------
@app.route('/tarjar_docx', methods=['GET', 'POST'])
def tarjar_docx_preview():
    if request.method == 'POST':
        arquivo = request.files.get("docxfile")
        selecionados = request.form.getlist("itens")

        if not arquivo or not arquivo.filename.endswith('.docx'):
            return "Arquivo inválido. Envie um .docx.", 400

        padroes_ativos = {k: v for k, v in PADROES_SENSIVEIS.items() if k in selecionados}

        conteudo_bytes = arquivo.read()
        file_stream = io.BytesIO(conteudo_bytes)
        doc = Document(file_stream)

        ocorrencias = []
        paragrafos_com_tarja = []

        for i, par in enumerate(doc.paragraphs):
            texto = par.text
            texto_tarjado = texto  # manter original para sobrescrever com tarjas
            offset = 0  # controle de deslocamento conforme o texto é alterado

            for tipo, regex in padroes_ativos.items():
                for m in re.finditer(regex, texto):
                    encontrado = m.group()
                    inicio = m.start() + offset
                    fim = m.end() + offset
                    tarja = '█' * len(encontrado)
                    texto_tarjado = (
                        texto_tarjado[:inicio] + tarja + texto_tarjado[fim:]
                    )

                    # Atualiza o offset após substituir
                    offset += len(tarja) - len(encontrado)

                    ocorrencias.append({
                        "tipo": tipo,
                        "texto": encontrado,
                        "paragrafo": i,
                        "start": m.start(),
                        "end": m.end(),
                        "id": f"{i}_{m.start()}_{m.end()}"
                    })

            paragrafos_com_tarja.append(texto_tarjado)

        # Salva cópia temporária do original para edição posterior
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        with open(temp_path, "wb") as f:
            f.write(conteudo_bytes)

        session['doc_ocorrencias'] = ocorrencias
        session['doc_path'] = temp_path

        return render_template("preview_docx.html", ocorrencias=ocorrencias, paragrafos=paragrafos_com_tarja)

    return render_template("tarjar_docx.html", padroes=PADROES_SENSIVEIS.keys())

@app.route("/aplicar_tarjas_docx", methods=["POST"])
def aplicar_tarjas_docx():
    selecionados = request.form.getlist("selecionados")
    trechos_manuais_raw = request.form.get("tarjas_manualmente_adicionadas", "")
    trechos_manuais = [t.strip() for t in trechos_manuais_raw.split("|") if t.strip()]

    ocorrencias = session.get("doc_ocorrencias", [])
    caminho = session.get("doc_path", None)

    if not caminho or not os.path.exists(caminho):
        return "Erro: Arquivo temporário não encontrado.", 400

    doc = Document(caminho)

    paragrafo_edits = {}

    for item in ocorrencias:
        if item["id"] in selecionados:
            idx = item["paragrafo"]
            texto_original = doc.paragraphs[idx].text
            if idx not in paragrafo_edits:
                paragrafo_edits[idx] = texto_original

            start, end = item["start"], item["end"]
            trecho = texto_original[start:end]
            texto_editado = paragrafo_edits[idx].replace(trecho, "█" * len(trecho), 1)
            paragrafo_edits[idx] = texto_editado

    if trechos_manuais:
        for i, par in enumerate(doc.paragraphs):
            texto = paragrafo_edits.get(i, par.text)
            for trecho_manual in trechos_manuais:
                if trecho_manual in texto:
                    texto = texto.replace(trecho_manual, "█" * len(trecho_manual))
                    paragrafo_edits[i] = texto

    for i, novo_texto in paragrafo_edits.items():
        par = doc.paragraphs[i]
        par.clear()
        run = par.add_run(novo_texto)
        run.font.color.rgb = RGBColor(0, 0, 0)

    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)

    # NÃO REMOVER o arquivo aqui! Senão o send_file não encontra.
    # os.remove(caminho)

    return send_file(
        mem_file,
        as_attachment=True,
        download_name="documento_tarjado.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/atualizar_preview_docx", methods=["POST"])
def atualizar_preview_docx():
    try:
        # Recebe JSON do frontend
        data = request.get_json(force=True)
        selecionados = set(data.get("selecionados", []))
        trechos_manuais = data.get("manuais", [])

        # Obtém ocorrências automáticas e caminho do DOCX da sessão
        ocorrencias = session.get("doc_ocorrencias", [])
        caminho = session.get("doc_path", None)

        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400

        doc = Document(caminho)
        paragrafo_edits = {}

        # --- Tarjas automáticas selecionadas ---
        for item in ocorrencias:
            if item["id"] in selecionados:
                idx = item["paragrafo"]
                texto_original = doc.paragraphs[idx].text
                if idx not in paragrafo_edits:
                    paragrafo_edits[idx] = texto_original
                start, end = item["start"], item["end"]
                trecho = texto_original[start:end]
                # Substituição com tarja
                texto_editado = paragrafo_edits[idx][:start] + "█" * len(trecho) + paragrafo_edits[idx][end:]
                paragrafo_edits[idx] = texto_editado

        # --- Tarjas manuais ---
        for i, par in enumerate(doc.paragraphs):
            texto = paragrafo_edits.get(i, par.text)
            for trecho_manual in trechos_manuais:
                # Substituição case-insensitive
                texto = re.sub(re.escape(trecho_manual), lambda m: "█" * len(m.group()), texto, flags=re.IGNORECASE)
            paragrafo_edits[i] = texto  # garante atualização final

        # Monta lista de parágrafos atualizados
        paragrafos_atualizados = [paragrafo_edits.get(i, doc.paragraphs[i].text) for i in range(len(doc.paragraphs))]

        return jsonify({"paragrafos": paragrafos_atualizados})

    except Exception as e:
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

# ----------------------------------------------------------------------------------- Padrões para PDF ----------------------------------------------------------------------------------
@app.route('/tarjar_pdf', methods=['GET', 'POST'])
def tarjar_pdf():
    if request.method == 'POST':
        arquivo = request.files.get('pdffile')
        tipos_selecionados = request.form.getlist('tipos')  # incluir 'qrcode' no front-end
        print("DEBUG tipos_selecionados:", tipos_selecionados)

        if not arquivo or not arquivo.filename.endswith('.pdf'):
            return "Arquivo inválido. Envie um .pdf.", 400

        padroes_filtrados = {k: v for k, v in PADROES_SENSIVEIS.items() if k in tipos_selecionados}

        pdf_bytes = arquivo.read()
        doc = fitz.open("pdf", pdf_bytes)

        ocorrencias = []
        redactions_por_pagina = {}

        for pagina_num in range(len(doc)):
            pagina = doc[pagina_num]
            texto = pagina.get_text("text")

            # EXISTENTE: detecção por regex / texto
            for tipo, regex in padroes_filtrados.items():
                for m in re.finditer(regex, texto):
                    termo = m.group()
                    ocorrencias.append({
                        "id": f"{pagina_num}_{m.start()}_{m.end()}",
                        "tipo": tipo,
                        "texto": termo,
                        "pagina": pagina_num,
                        "start": m.start(),
                        "end": m.end()
                    })

                    areas = pagina.search_for(termo)
                    for area in areas:
                        redactions_por_pagina.setdefault(pagina_num, []).append(area)

            # NOVO: detecção de QR Codes (se marcado no front)
            if 'qrcode' in tipos_selecionados:
                qrs = detectar_qrcodes_pagina(pagina, pagina_num, dpi=200)
                print(f"DEBUG página {pagina_num+1}: detectados {len(qrs)} QR(s)")
                for qr in qrs:
                    print("DEBUG bbox:", qr.get("bbox"))
                    ocorrencias.append(qr)
                    rect = fitz.Rect(*qr['bbox'])
                    redactions_por_pagina.setdefault(pagina_num, []).append(rect)

            # OPCIONAL: modo teste (pinta um retângulo fixo no canto inf. direito)
            if 'qrcode_fixed' in tipos_selecionados:
                try:
                    pagina_rect = pagina.rect
                    largura = 128  # ~45mm
                    altura  = 128
                    x1 = pagina_rect.x1 - 30  # margem direita 30pt
                    y1 = 30                    # margem inferior 30pt
                    rect_fix = fitz.Rect(x1 - largura, y1, x1, y1 + altura)
                    redactions_por_pagina.setdefault(pagina_num, []).append(rect_fix)
                    ocorrencias.append({
                        "id": f"qrf_{pagina_num}_0",
                        "tipo": "qrcode_fixed",
                        "texto": "QR fixo (teste)",
                        "pagina": pagina_num,
                        "bbox": [rect_fix.x0, rect_fix.y0, rect_fix.x1, rect_fix.y1]
                    })
                    print(f"DEBUG aplicado qrcode_fixed na página {pagina_num+1}")
                except Exception as e:
                    print("DEBUG erro no qrcode_fixed:", e)

        # Aplicar redactions (apenas para visualização/preview)
        for pagina_idx, areas in redactions_por_pagina.items():
            pagina = doc[pagina_idx]
            for area in areas:
                pagina.add_redact_annot(area, fill=(0, 0, 0))
            pagina.apply_redactions()

        # Salvar o PDF modificado em memória
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        doc.close()

        pdf_b64 = base64.b64encode(mem_file.read()).decode('utf-8')

        # Salvar original temp (para aplicar tarjas reais depois)
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        with open(temp_path, 'wb') as f:
            f.write(pdf_bytes)

        session['pdf_path'] = temp_path
        session['pdf_ocorrencias'] = ocorrencias

        return render_template("preview_pdf.html", ocorrencias=ocorrencias, pdf_data=pdf_b64)

    return render_template('tarjar_pdf.html', padroes=PADROES_SENSIVEIS.keys())

@app.route('/aplicar_tarjas_pdf', methods=['POST'])
def aplicar_tarjas_pdf():
    selecionados = request.form.getlist('selecionados')
    preservar_logo = request.form.get('preservar_logo', '0') == '1'

    trechos_manuais_raw = request.form.get('tarjas_manualmente_adicionadas', '')
    trechos_manuais = [t.strip() for t in trechos_manuais_raw.split('|') if t.strip()]

    # NOVO: flags vindas do preview
    qrcode_preview = request.form.get('qrcode_preview', '0') == '1'
    qrcode_fixed_preview = request.form.get('qrcode_fixed_preview', '0') == '1'

    caminho = session.get('pdf_path')
    ocorrencias = session.get('pdf_ocorrencias', [])

    if not caminho or not os.path.exists(caminho):
        return "Erro: arquivo temporário não encontrado.", 400

    doc = fitz.open(caminho)
    redactions_por_pagina = {}

    # 1) Ocorrências automáticas selecionadas (texto / bbox)
    for item in ocorrencias:
        if item['id'] in selecionados:
            pagina_idx = item['pagina']
            pagina = doc[pagina_idx]
            if 'bbox' in item:
                area = fitz.Rect(*item['bbox'])
                if not (preservar_logo and area.y0 < 100):
                    redactions_por_pagina.setdefault(pagina_idx, []).append(area)
            else:
                termo = item['texto']
                areas = pagina.search_for(termo)
                for area in areas:
                    if preservar_logo and area.y0 < 100:
                        continue
                    redactions_por_pagina.setdefault(pagina_idx, []).append(area)

    # 2) Tarjas manuais por texto
    if trechos_manuais:
        for num_pagina in range(len(doc)):
            pagina = doc[num_pagina]
            texto_pagina = pagina.get_text()
            for trecho in trechos_manuais:
                if trecho in texto_pagina:
                    areas = pagina.search_for(trecho)
                    for area in areas:
                        if preservar_logo and area.y0 < 100:
                            continue
                        redactions_por_pagina.setdefault(num_pagina, []).append(area)

    # 3) NOVO: QR Code (detecção) conforme o estado do preview
    if qrcode_preview:
        for num_pagina in range(len(doc)):
            pagina = doc[num_pagina]
            try:
                qrs = detectar_qrcodes_pagina(pagina, num_pagina, dpi=200)
            except Exception as e:
                qrs = []
                print("DEBUG detectar_qrcodes_pagina falhou no apply:", e)
            for qr in qrs:
                area = fitz.Rect(*qr['bbox'])
                if not (preservar_logo and area.y0 < 100):
                    redactions_por_pagina.setdefault(num_pagina, []).append(area)

    # 4) NOVO: QR fixo (teste) conforme o estado do preview
    if qrcode_fixed_preview:
        for num_pagina in range(len(doc)):
            pagina = doc[num_pagina]
            try:
                pagina_rect = pagina.rect
                largura = 128
                altura = 128
                x1 = pagina_rect.x1 - 30
                y1 = 30
                rect_fix = fitz.Rect(x1 - largura, y1, x1, y1 + altura)
                if not (preservar_logo and rect_fix.y0 < 100):
                    redactions_por_pagina.setdefault(num_pagina, []).append(rect_fix)
            except Exception as e:
                print("DEBUG qrcode_fixed apply falhou:", e)

    # Aplicar e retornar
    for pagina_idx, areas in redactions_por_pagina.items():
        pagina = doc[pagina_idx]
        for area in areas:
            pagina.add_redact_annot(area, fill=(0, 0, 0))
        pagina.apply_redactions()

    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)
    doc.close()

    return send_file(
        mem_file,
        as_attachment=True,
        download_name="documento_tarjado.pdf",
        mimetype="application/pdf"
    )

@app.route('/preview_pdf', methods=['POST'])
def preview_pdf():
    arquivo = request.files['arquivo']
    nome_temporario = os.path.join('uploads', f"{uuid.uuid4()}.pdf")
    arquivo.save(nome_temporario)

    doc = fitz.open(nome_temporario)

    pdf_data = base64.b64encode(open(nome_temporario, "rb").read()).decode('utf-8')

    ocorrencias = detectar_dados(doc)  # sua função atual
    texto_extraido = ""
    for pagina in doc:
        texto_extraido += pagina.get_text() + "\n"

    doc.close()

    session['pdf_path'] = nome_temporario
    session['pdf_ocorrencias'] = ocorrencias

    return render_template(
        "preview_pdf.html",
        pdf_data=pdf_data,
        ocorrencias=ocorrencias,
        texto_extraido=texto_extraido
    )

@app.route('/atualizar_preview_pdf', methods=['POST'])
def atualizar_preview_pdf():
    try:
        data = request.get_json(force=True)
        selecionados = data.get("selecionados", [])
        trechos_manuais = data.get("manuais", [])

        # NOVO: flags vindas do preview
        qrcode = bool(data.get("qrcode", False))
        qrcode_fixed = bool(data.get("qrcode_fixed", False))

        caminho = session.get('pdf_path')
        ocorrencias = session.get('pdf_ocorrencias', [])

        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400

        doc = fitz.open(caminho)
        redactions_por_pagina = {}

        # 1) Ocorrências automáticas selecionadas (texto e itens com bbox)
        for item in ocorrencias:
            if item['id'] in selecionados:
                pagina_idx = item['pagina']
                pagina = doc[pagina_idx]
                if 'bbox' in item:
                    area = fitz.Rect(*item['bbox'])
                    redactions_por_pagina.setdefault(pagina_idx, []).append(area)
                else:
                    termo = item['texto']
                    areas = pagina.search_for(termo)
                    for area in areas:
                        redactions_por_pagina.setdefault(pagina_idx, []).append(area)

        # 2) Tarjas manuais por texto
        for num_pagina in range(len(doc)):
            pagina = doc[num_pagina]
            texto_pagina = pagina.get_text()
            for trecho in trechos_manuais:
                if trecho and trecho in texto_pagina:
                    areas = pagina.search_for(trecho)
                    for area in areas:
                        redactions_por_pagina.setdefault(num_pagina, []).append(area)

        # 3) NOVO: QR Code na PRÉ-VISUALIZAÇÃO
        if qrcode:
            for num_pagina in range(len(doc)):
                pagina = doc[num_pagina]
                try:
                    qrs = detectar_qrcodes_pagina(pagina, num_pagina, dpi=200)
                except Exception as e:
                    qrs = []
                    print("DEBUG detectar_qrcodes_pagina falhou no preview:", e)
                for qr in qrs:
                    area = fitz.Rect(*qr['bbox'])
                    redactions_por_pagina.setdefault(num_pagina, []).append(area)

        # 4) NOVO: QR fixo (teste) na PRÉ-VISUALIZAÇÃO
        if qrcode_fixed:
            for num_pagina in range(len(doc)):
                pagina = doc[num_pagina]
                try:
                    pagina_rect = pagina.rect
                    largura = 128  # ~45 mm
                    altura = 128
                    x1 = pagina_rect.x1 - 30  # margem direita 30pt
                    y1 = 30                    # margem inferior 30pt
                    rect_fix = fitz.Rect(x1 - largura, y1, x1, y1 + altura)
                    redactions_por_pagina.setdefault(num_pagina, []).append(rect_fix)
                except Exception as e:
                    print("DEBUG qrcode_fixed preview falhou:", e)

        # Aplicar redactions e retornar base64
        for pagina_idx, areas in redactions_por_pagina.items():
            pagina = doc[pagina_idx]
            for area in areas:
                pagina.add_redact_annot(area, fill=(0, 0, 0))
            pagina.apply_redactions()

        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        doc.close()

        pdf_b64 = base64.b64encode(mem_file.read()).decode('utf-8')
        return jsonify({"pdf_data": pdf_b64})

    except Exception as e:
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

@app.route('/download_pdf_tarjado')
def download_pdf_tarjado():
    path = session.get('pdf_tarjado_path', None)
    if not path or not os.path.exists(path):
        return "Nenhum PDF tarjado disponível.", 400

    return send_file(path, as_attachment=True, download_name="documento_tarjado.pdf", mimetype="application/pdf")

# ----------------------------------------------------------------------------------- Padrões para PDF OCR ----------------------------------------------------------------------------------
@app.route('/tarjar_ocr_pdf', methods=['GET', 'POST'])
def tarjar_ocr_pdf():
    if request.method == 'POST':
        arquivo = request.files.get('ocrpdf')
        tipos_selecionados = request.form.getlist('tipos')

        if not arquivo or not arquivo.filename.lower().endswith('.pdf'):
            return "Arquivo inválido. Envie um arquivo PDF escaneado.", 400

        padroes_ativos = {k: v for k, v in PADROES_SENSIVEIS.items() if k in tipos_selecionados}

        try:
            pdf_bytes = arquivo.read()
            imagens = convert_from_bytes(pdf_bytes)
        except Exception as e:
            app.logger.error(f"Erro ao converter PDF em imagens: {e}")
            return "Erro ao processar o arquivo PDF.", 500

        # ⚠️ NÃO TARJAR AQUI – só coletar ocorrências!
        todas_ocorrencias = []

        for idx, imagem in enumerate(imagens):
            dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)

            for tipo, regex in padroes_ativos.items():
                try:
                    pattern = regex if isinstance(regex, re.Pattern) else re.compile(regex, re.IGNORECASE | re.UNICODE)
                except re.error as e:
                    app.logger.error(f"Regex inválido para tipo '{tipo}': {e}")
                    continue

                for i, palavra in enumerate(dados_ocr['text']):
                    texto = (palavra or '').strip()
                    if not texto:
                        continue
                    if pattern.search(texto):
                        todas_ocorrencias.append({
                            "id": str(uuid.uuid4()),
                            "pagina": idx,
                            "tipo": tipo,
                            "texto": texto
                        })

        # --- BUSCA MANUAL POR TEXTO DIGITADO (só coletar) ---
        texto_manual = request.form.get('tarjas_manualmente_adicionadas', '').strip()
        if texto_manual:
            trechos_manualmente_adicionados = [t.strip() for t in texto_manual.split('|') if t.strip()]
            for idx, imagem in enumerate(imagens):
                dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)
                # Vamos apenas registrar a ocorrência “manual”; a tarja será desenhada na rota de preview/download
                for trecho in trechos_manualmente_adicionados:
                    todas_ocorrencias.append({
                        "id": str(uuid.uuid4()),
                        "pagina": idx,
                        "tipo": "manual",
                        "texto": trecho
                    })

        # ✅ Salvar o PDF ORIGINAL (sem tarjas) para ser a base do preview e do download
        diretorio_temp = os.path.join(app.root_path, 'arquivos_temp')
        os.makedirs(diretorio_temp, exist_ok=True)
        nome_arquivo = f"{uuid.uuid4()}_original.pdf"
        caminho_arquivo = os.path.join(diretorio_temp, nome_arquivo)
        with open(caminho_arquivo, 'wb') as f:
            f.write(pdf_bytes)

        # Guardar em sessão
        session['ocr_original_pdf_path'] = caminho_arquivo
        session['ocr_ocorrencias'] = todas_ocorrencias

        # Renderizar preview (o JS já chama /atualizar_preview_ocr_pdf na carga)
        return render_template(
            "preview_ocr.html",
            ocorrencias=todas_ocorrencias
        )

    return render_template('tarjar_ocr_pdf.html', padroes=PADROES_SENSIVEIS.keys())

@app.route('/aplicar_tarjas_ocr_pdf', methods=['POST'])
def aplicar_tarjas_ocr_pdf():
    caminho = session.get('ocr_original_pdf_path')  # <<<< usar o ORIGINAL
    ocorrencias_automaticas = session.get('ocr_ocorrencias', [])

    if not caminho or not os.path.exists(caminho):
        return "Arquivo OCR não encontrado.", 400

    imagens = convert_from_bytes(open(caminho, 'rb').read())
    imagens_tarjadas = [img.copy() for img in imagens]

    selecionados = request.form.getlist('selecionados')
    selecionados_set = set(str(s) for s in selecionados)

    # Campo oculto unificado (pipe-separado)
    texto_manual = request.form.get('tarjas_manualmente_adicionadas', '').strip()
    trechos_manuais = [t.strip().lower() for t in texto_manual.split('|') if t.strip()]

    tarjas_aplicadas = []

    for idx, imagem in enumerate(imagens_tarjadas):
        draw = ImageDraw.Draw(imagem)
        dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)

        linhas = {}
        for i in range(len(dados_ocr['text'])):
            linha_num = dados_ocr['line_num'][i]
            if linha_num not in linhas:
                linhas[linha_num] = []
            linhas[linha_num].append({
                'text': (dados_ocr['text'][i] or '').strip(),
                'left': int(dados_ocr['left'][i]),
                'top': int(dados_ocr['top'][i]),
                'width': int(dados_ocr['width'][i]),
                'height': int(dados_ocr['height'][i])
            })

        # Selecionados automáticos
        for ocorrencia in [o for o in ocorrencias_automaticas if o['pagina'] == idx]:
            if str(ocorrencia['id']) not in selecionados_set:
                continue
            termo_lower = ocorrencia['texto'].lower()
            for palavras_linha in linhas.values():
                linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                if termo_lower in linha_texto:
                    char_count = 0
                    trecho_start = linha_texto.find(termo_lower)
                    trecho_end = trecho_start + len(termo_lower)
                    for palavra in palavras_linha:
                        palavra_start = char_count
                        palavra_end = char_count + len(palavra['text'])
                        char_count += len(palavra['text']) + 1
                        if palavra_end > trecho_start and palavra_start < trecho_end:
                            x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                            draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                            tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})

        # Manuais
        for trecho in trechos_manuais:
            trecho_lower = trecho.lower()
            for palavras_linha in linhas.values():
                linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                if fuzz.partial_ratio(trecho_lower, linha_texto) >= 85:
                    char_count = 0
                    trecho_start = linha_texto.find(trecho_lower)
                    trecho_end = trecho_start + len(trecho_lower)
                    for palavra in palavras_linha:
                        palavra_start = char_count
                        palavra_end = char_count + len(palavra['text'])
                        char_count += len(palavra['text']) + 1
                        if palavra_end > trecho_start and palavra_start < trecho_end:
                            x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                            draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                            tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})

    session['tarjas_ocr'] = tarjas_aplicadas

    buffer = io.BytesIO()
    imagens_tarjadas[0].save(buffer, format="PDF", save_all=True, append_images=imagens_tarjadas[1:])
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="documento_tarjado.pdf",
        mimetype="application/pdf"
    )

@app.route('/atualizar_preview_ocr_pdf', methods=['POST'])
def atualizar_preview_ocr_pdf():
    try:
        data = request.get_json(force=True)
        selecionados = data.get("selecionados", [])
        trechos_manuais = data.get("manuais", [])
        selecionados_set = set(str(s) for s in selecionados)

        caminho = session.get('ocr_original_pdf_path')  # <<<< usar o ORIGINAL
        ocorrencias = session.get('ocr_ocorrencias', [])

        if not caminho or not os.path.exists(caminho):
            return jsonify({"erro": "Arquivo temporário não encontrado."}), 400

        imagens = convert_from_bytes(open(caminho, 'rb').read())
        imagens_tarjadas = [img.copy() for img in imagens]
        tarjas_aplicadas = []

        for idx, imagem in enumerate(imagens_tarjadas):
            draw = ImageDraw.Draw(imagem)
            dados_ocr = pytesseract.image_to_data(imagem, lang='por', output_type=pytesseract.Output.DICT)

            # Agrupar por linha (igual ao seu código)
            linhas = {}
            for i in range(len(dados_ocr['text'])):
                linha_num = dados_ocr['line_num'][i]
                if linha_num not in linhas:
                    linhas[linha_num] = []
                linhas[linha_num].append({
                    'text': (dados_ocr['text'][i] or '').strip(),
                    'left': int(dados_ocr['left'][i]),
                    'top': int(dados_ocr['top'][i]),
                    'width': int(dados_ocr['width'][i]),
                    'height': int(dados_ocr['height'][i])
                })

            # Tarjas automáticas SELECIONADAS
            for ocorrencia in [o for o in ocorrencias if o['pagina'] == idx]:
                if str(ocorrencia['id']) not in selecionados_set:
                    continue
                termo_lower = ocorrencia['texto'].lower()
                for palavras_linha in linhas.values():
                    linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                    if termo_lower in linha_texto:
                        char_count = 0
                        trecho_start = linha_texto.find(termo_lower)
                        trecho_end = trecho_start + len(termo_lower)
                        for palavra in palavras_linha:
                            palavra_start = char_count
                            palavra_end = char_count + len(palavra['text'])
                            char_count += len(palavra['text']) + 1
                            if palavra_end > trecho_start and palavra_start < trecho_end:
                                x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                                draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                                tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})

            # Tarjas manuais (texto livre do preview)
            for trecho in trechos_manuais:
                trecho_lower = trecho.lower()
                for palavras_linha in linhas.values():
                    linha_texto = ' '.join([p['text'] for p in palavras_linha]).lower()
                    if fuzz.partial_ratio(trecho_lower, linha_texto) >= 85:
                        char_count = 0
                        trecho_start = linha_texto.find(trecho_lower)
                        trecho_end = trecho_start + len(trecho_lower)
                        for palavra in palavras_linha:
                            palavra_start = char_count
                            palavra_end = char_count + len(palavra['text'])
                            char_count += len(palavra['text']) + 1
                            if palavra_end > trecho_start and palavra_start < trecho_end:
                                x, y, w, h = palavra['left'], palavra['top'], palavra['width'], palavra['height']
                                draw.rectangle([(x, y), (x + w, y + h)], fill='black')
                                tarjas_aplicadas.append({'pagina': idx, 'texto': palavra['text']})

        session['tarjas_ocr'] = tarjas_aplicadas

        pdf_mem = io.BytesIO()
        imagens_tarjadas[0].save(pdf_mem, format="PDF", save_all=True, append_images=imagens_tarjadas[1:])
        pdf_mem.seek(0)
        pdf_b64 = base64.b64encode(pdf_mem.read()).decode('utf-8')

        return jsonify({"pdf_data": pdf_b64})

    except Exception as e:
        app.logger.error(f"Erro ao atualizar preview OCR PDF: {e}")
        return jsonify({"erro": f"Erro no servidor: {str(e)}"}), 500

@app.route('/download_pdf_ocr')
def download_pdf_ocr():
    caminho = session.get('ocr_pdf_path')

    if not caminho or not os.path.exists(caminho):
        return "Arquivo não encontrado.", 404

    return send_file(caminho, as_attachment=True, download_name="pdf_tarjado_ocr.pdf")

@app.route('/ver_pdf_ocr')
def ver_pdf_ocr():
    caminho = session.get('ocr_pdf_path')
    if not caminho or not os.path.exists(caminho):
        return "Arquivo não encontrado.", 404

    return send_file(caminho, mimetype='application/pdf')

# ----------------------------------------------------------------------------------- Padrões para qrCode ----------------------------------------------------------------------------------
def detectar_qrcodes_pagina(pagina, pagina_num, dpi=200):
    """
    Renderiza a página em imagem e usa pyzbar/ZBar para localizar QR Codes.
    Retorna lista de dicts com id, tipo='qrcode', pagina e bbox [x0, y0, x1, y1].
    """
    if not _ZBAR_OK:
        print("DEBUG ZBar não disponível (_ZBAR_OK=False).")
        return []

    # Render da página -> imagem PIL
    try:
        pix = pagina.get_pixmap(dpi=dpi, alpha=False)
        mode = "RGB" if not pix.alpha else "RGBA"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
    except Exception as e:
        print("DEBUG erro ao renderizar página:", e)
        return []

    # Detecta com ZBar
    try:
        resultados = zbar_decode(img)
    except Exception as e:
        print("DEBUG zbar_decode falhou (falta DLL do ZBar?):", e)
        return []

    if not resultados:
        return []

    # pixels -> pontos PDF (1 pt = 1/72in; render em dpi)
    escala = dpi / 72.0
    encontrados = []
    for i, r in enumerate(resultados):
        if getattr(r, "type", "") != "QRCODE":
            continue

        left, top, w, h = r.rect.left, r.rect.top, r.rect.width, r.rect.height
        x0 = left / escala
        y0 = top / escala
        x1 = (left + w) / escala
        y1 = (top + h) / escala

        texto_lido = ""
        try:
            texto_lido = (r.data or b"").decode("utf-8", "ignore")
        except Exception:
            pass

        encontrados.append({
            "id": f"qr_{pagina_num}_{i}",
            "tipo": "qrcode",
            "texto": texto_lido if texto_lido else "QR Code",
            "pagina": pagina_num,
            "bbox": [x0, y0, x1, y1],
        })
    return encontrados