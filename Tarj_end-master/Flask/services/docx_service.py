import io
import os
import re
import tempfile
from docx import Document
from docx.shared import RGBColor

import spacy
nlp = spacy.load("pt_core_news_md")  

def encontrar_ocorrencias_docx(conteudo_bytes, padroes_ativos, usar_spacy=False):
    file_stream = io.BytesIO(conteudo_bytes)
    doc = Document(file_stream)

    ocorrencias = []
    paragrafos_com_tarja = []

    for i, par in enumerate(doc.paragraphs):
        texto = par.text
        texto_tarjado = texto
        offset = 0

        # Sempre aplica Regex
        for tipo, regex in padroes_ativos.items():
            for m in re.finditer(regex, texto):
                encontrado = m.group()
                inicio = m.start() + offset
                fim = m.end() + offset
                tarja = '█' * len(encontrado)
                texto_tarjado = texto_tarjado[:inicio] + tarja + texto_tarjado[fim:]
                offset += len(tarja) - len(encontrado)

                ocorrencias.append({
                    "tipo": tipo,
                    "texto": encontrado,
                    "paragrafo": i,
                    "start": m.start(),
                    "end": m.end(),
                    "id": f"{i}_{m.start()}_{m.end()}"
                })

        # Opcional: aplica SpaCy
        if usar_spacy:
            doc_spacy = nlp(texto)
            for ent in doc_spacy.ents:
                encontrado = ent.text
                inicio, fim = ent.start_char, ent.end_char
                tarja = '█' * len(encontrado)
                texto_tarjado = texto_tarjado[:inicio] + tarja + texto_tarjado[fim:]

                ocorrencias.append({
                    "tipo": ent.label_,
                    "texto": encontrado,
                    "paragrafo": i,
                    "start": inicio,
                    "end": fim,
                    "id": f"{i}_{inicio}_{fim}"
                })

        paragrafos_com_tarja.append(texto_tarjado)

    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    with open(temp_path, "wb") as f:
        f.write(conteudo_bytes)

    return ocorrencias, paragrafos_com_tarja, temp_path

def aplicar_tarjas_docx(caminho, ocorrencias, selecionados, trechos_manuais):
    doc = Document(caminho)
    paragrafo_edits = {}

    for item in ocorrencias:
        if item["id"] in selecionados:
            idx = item["paragrafo"]
            texto_original = doc.paragraphs[idx].text
            if idx not in paragrafo_edits:
                paragrafo_edits[idx] = texto_original

            start, end = item["start"], item["end"]
            trecho = texto_original[start:end]
            texto_editado = paragrafo_edits[idx].replace(trecho, "█" * len(trecho), 1)
            paragrafo_edits[idx] = texto_editado

    if trechos_manuais:
        for i, par in enumerate(doc.paragraphs):
            texto = paragrafo_edits.get(i, par.text)
            for trecho_manual in trechos_manuais:
                texto = texto.replace(trecho_manual, "█" * len(trecho_manual))
                paragrafo_edits[i] = texto

    for i, novo_texto in paragrafo_edits.items():
        par = doc.paragraphs[i]
        par.clear()
        run = par.add_run(novo_texto)
        run.font.color.rgb = RGBColor(0, 0, 0)

    mem_file = io.BytesIO()
    doc.save(mem_file)
    mem_file.seek(0)

    return mem_file


def atualizar_preview_docx(caminho, ocorrencias, selecionados, trechos_manuais, usar_spacy=False):
    doc = Document(caminho)
    paragrafo_edits = {}

    # Sempre aplica regex baseado nas ocorrências salvas
    for item in ocorrencias:
        if item["id"] in selecionados:
            idx = item["paragrafo"]
            texto_original = doc.paragraphs[idx].text
            if idx not in paragrafo_edits:
                paragrafo_edits[idx] = texto_original
            start, end = item["start"], item["end"]
            trecho = texto_original[start:end]
            texto_editado = paragrafo_edits[idx][:start] + "█" * len(trecho) + paragrafo_edits[idx][end:]
            paragrafo_edits[idx] = texto_editado

    # Aplica spaCy se habilitado
    if usar_spacy:
        for i, par in enumerate(doc.paragraphs):
            texto = paragrafo_edits.get(i, par.text)
            doc_spacy = nlp(texto)
            for ent in doc_spacy.ents:
                trecho = ent.text
                texto = texto.replace(trecho, "█" * len(trecho))
            paragrafo_edits[i] = texto

    # Aplica trechos manuais
    for i, par in enumerate(doc.paragraphs):
        texto = paragrafo_edits.get(i, par.text)
        for trecho_manual in trechos_manuais:
            texto = re.sub(re.escape(trecho_manual), lambda m: "█" * len(m.group()), texto, flags=re.IGNORECASE)
        paragrafo_edits[i] = texto

    return [paragrafo_edits.get(i, doc.paragraphs[i].text) for i in range(len(doc.paragraphs))]
